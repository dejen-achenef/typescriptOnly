#!/usr/bin/env python3
"""
Script to create a DOCX template for the document scanner app.
This template will be used by docx_template package to insert scanned images.

Requirements:
    pip install python-docx

Usage:
    python scripts/create_docx_template.py
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_template():
    """Create a DOCX template with content controls for images."""
    doc = Document()
    
    # Set up page margins (narrow margins for scanned documents)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add a paragraph with template instructions (will be replaced)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add placeholder text for the template
    # The docx_template library will look for {{pages}} and {{img}} tags
    run = p.add_run('{% for page in pages %}')
    run.font.size = Pt(1)  # Make it tiny so it's not visible
    
    # Add image placeholder
    # In docx_template, we use {{%img}} for image insertion
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('{{%img}}')
    run2.font.size = Pt(1)
    
    # Add page break placeholder
    p3 = doc.add_paragraph()
    run3 = p3.add_run('{% endfor %}')
    run3.font.size = Pt(1)
    
    # Save the template
    output_path = '../assets/docx/document_template.docx'
    doc.save(output_path)
    print(f'✓ Template created successfully at: {output_path}')
    print('  This template can now be used by the Flutter app to generate DOCX files.')

if __name__ == '__main__':
    try:
        create_template()
    except ImportError:
        print('Error: python-docx is not installed.')
        print('Install it with: pip install python-docx')
        exit(1)
    except Exception as e:
        print(f'Error creating template: {e}')
        exit(1)
